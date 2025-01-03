# -*- coding: utf-8 -*-
# Part of Odoo. See LICENSE file for full copyright and licensing details.
import base64
import docx
import io
from datetime import date, timedelta
from dateutil.relativedelta import relativedelta
from num2words import num2words
from odoo.exceptions import ValidationError

from odoo import models, fields, api, _

ordinal_suffix_days = {1: 'st', 2: 'nd', 3: 'rd'}


class EmployeeLetterWizards(models.TransientModel):
    _name = 'employee.letter.wizard'
    _description = "Employee Letter Wizard"

    def _letter(self):
        if 'employee_id' in self.env.context:
            employee_obj = self.env['hr.employee'].browse(self.env.context.get('employee_id'))
            if 'appraisal_id' in self.env.context:
                doc_type = self.env['employee.document.type'].search([('name', '=', 'Appraisal')])
                return [('allowed_in_stages.field_name', 'in', [employee_obj.emp_state]),
                        ('document_type_id', '=', doc_type.id)]
            else:
                doc_type = self.env['employee.document.type'].search([('name', 'in', ['Appraisal', 'Termination'])])
                return [('allowed_in_stages.field_name', 'in', [employee_obj.emp_state]),
                        ('document_type_id', 'not in', doc_type.ids)]
        else:
            return None

    letter = fields.Many2one('hr.letters', string="Letter Name", domain=_letter)

    @api.model
    def _change_content(self, paragraph, employee_obj):
        result = ""
        i = 0

        while i < len(paragraph):
            if paragraph[i:i + 2] == "{{":
                # Find the closing "}}"
                end_index = paragraph.find("}}", i + 2)

                if end_index != -1:
                    # Extract the content between "{{" and "}}"
                    variable_content = paragraph[i + 2:end_index].split(" - ")

                    # Replace the "{{...}}" with it Object Value
                    # Access the specified field of the record
                    if len(variable_content) == 1:
                        if variable_content[0].strip() == "emp_address":
                            replace_variable_content = f'{employee_obj.per_street}, {employee_obj.per_landmark}, {employee_obj.per_city}, {employee_obj.per_state_id.name}, {employee_obj.per_county_id.name}, {employee_obj.per_pcode}'
                        elif variable_content[0].strip() == "company_address":
                            replace_variable_content = f'{employee_obj.company_id.street}, {employee_obj.company_id.street2}, {employee_obj.company_id.city}, {employee_obj.company_id.state_id.name}, {employee_obj.company_id.country_id.name}, {employee_obj.company_id.zip}'
                        elif '%d' in variable_content[0].strip():
                            variable_content_1 = variable_content[0].split(".")
                            if 'today' in variable_content_1[0].strip():
                                if '%dth' in variable_content_1[1].strip():
                                    day = date.today().day
                                    day_suffix = ordinal_suffix_days.get(day,
                                                                         'th') if 10 <= day % 100 <= 20 else ordinal_suffix_days.get(
                                        day % 10, 'th')
                                    replace_variable_content = date.today().strftime(f'%d{day_suffix} %B, %Y')
                                else:
                                    replace_variable_content = date.today().strftime(variable_content_1[1].strip())
                            else:
                                pass
                        elif variable_content[0].strip() == "ms/mr/miss" or variable_content[0].strip() == "her/him" or \
                                variable_content[0].strip() == "her/his":
                            if employee_obj.gender == 'male':
                                pre_word_gender = ['Mr.', 'him', 'his']
                            elif employee_obj.gender == 'female':
                                if employee_obj.marital != 'married':
                                    pre_word_gender = ['Miss', 'her', 'her']
                                else:
                                    pre_word_gender = ['Ms.', 'her', 'her']
                            else:
                                pre_word_gender = ['', '', '']

                            if variable_content[0].strip() == "ms/mr/miss":
                                replace_variable_content = pre_word_gender[0]
                            elif variable_content[0].strip() == "her/him":
                                replace_variable_content = pre_word_gender[1]
                            else:
                                replace_variable_content = pre_word_gender[2]
                        elif variable_content[0].strip() == "sal_rev_dt":
                            emp_join_date = employee_obj.join_date
                            if emp_join_date.day > 15:
                                emp_join_date = emp_join_date.replace(day=1)
                                emp_join_date = emp_join_date + timedelta(days=32)
                                emp_join_date = emp_join_date.replace(day=1)
                            emp_join_date_day = emp_join_date.day
                            day_suffix = ordinal_suffix_days.get(emp_join_date_day,
                                                                 'th') if 10 <= emp_join_date_day % 100 <= 20 else ordinal_suffix_days.get(
                                emp_join_date_day % 10, 'th')
                            replace_variable_content = emp_join_date.strftime(f'%d{day_suffix} %B')
                        elif variable_content[0].strip() == "con_eff_dt":
                            emp_join_date = employee_obj.join_date
                            emp_join_date_day = emp_join_date.day
                            day_suffix = ordinal_suffix_days.get(emp_join_date_day,
                                                                 'th') if 10 <= emp_join_date_day % 100 <= 20 else ordinal_suffix_days.get(
                                emp_join_date_day % 10, 'th')

                            if emp_join_date.day > 15:
                                after_1_year_date = emp_join_date.replace(day=1)
                                after_1_year_date = after_1_year_date + relativedelta(years=1)
                                after_1_year_date = after_1_year_date + timedelta(days=32)
                                after_1_year_date = after_1_year_date.replace(day=1)
                                replace_variable_content = emp_join_date.strftime(
                                    f'%d{day_suffix} %B, %Y') + " to " + after_1_year_date.strftime(f'%dst %B, %Y')
                            else:
                                after_1_year_date = emp_join_date + relativedelta(years=1)
                                replace_variable_content = emp_join_date.strftime(
                                    f'%d{day_suffix} %B, %Y') + " to " + after_1_year_date.strftime(
                                    f'%d{day_suffix} %B, %Y')
                        elif variable_content[0].strip() == "con_app_dt":
                            emp_appraisal_date = employee_obj.appraisal_date

                            if emp_appraisal_date.day > 15:
                                after_1_year_date = emp_appraisal_date.replace(day=1)
                                after_1_year_date = after_1_year_date + relativedelta(years=1)
                                after_1_year_date = after_1_year_date + timedelta(days=32)
                                after_1_year_date = after_1_year_date.replace(day=1)
                                replace_variable_content = emp_appraisal_date.strftime(
                                    f'%B %Y') + " and terminating on " + after_1_year_date.strftime(f'%B %Y')
                            else:
                                after_1_year_date = emp_appraisal_date + relativedelta(years=1)
                                replace_variable_content = emp_appraisal_date.strftime(
                                    f'%B %Y') + " and terminating on " + after_1_year_date.strftime(
                                    f'%B %Y')
                        elif variable_content[0].strip() == "app_exis_year":
                            replace_variable_content = f'{str(employee_obj.appraisal_date.year - 1)} to {str(employee_obj.appraisal_date.year)}'
                        elif variable_content[0].strip() == "app_new_year":
                            replace_variable_content = f'{str(employee_obj.appraisal_date.year)} to {str(employee_obj.appraisal_date.year + 1)}'
                        elif 'salary' in variable_content[0].strip():
                            emp_latest_contract_obj = \
                                employee_obj.contract_ids.sorted(key=lambda x: x.create_date, reverse=True)[0]

                            # Yearly wage amount & Word ...
                            if variable_content[0].strip() == 'salary_annual_ctc_word':
                                yearly_wage_words = num2words(round(emp_latest_contract_obj.wage * 12), lang='en_IN')
                                yearly_wage_words = yearly_wage_words[0].capitalize() + yearly_wage_words[1:]
                                yearly_wage_amount_words = f'Rs. {round(emp_latest_contract_obj.wage * 12)}/- (Rupees {yearly_wage_words} Only)'
                                replace_variable_content = yearly_wage_amount_words

                            # Monthly wage amount & Word ...
                            elif variable_content[0].strip() == 'salary_monthly_ctc_word':
                                monthly_wage_words = num2words(round(emp_latest_contract_obj.wage), lang='en_IN')
                                monthly_wage_words = monthly_wage_words[0].capitalize() + monthly_wage_words[1:]
                                monthly_wage_amount_words = f'Rs. {round(emp_latest_contract_obj.wage)}/- (Rupees {monthly_wage_words} Only)'
                                replace_variable_content = monthly_wage_amount_words

                            else:
                                # Load Salary Component's ...
                                emp_latest_contract_obj.load_salary_components()

                                # Calculate Salary Component's
                                for salary_component in emp_latest_contract_obj.applicable_salary_rule_ids:
                                    if salary_component.rule_id.code == 'BASIC':
                                        monthly_basic = salary_component.amount
                                        yearly_basic = monthly_basic * 12
                                    elif salary_component.rule_id.code == 'CTC':
                                        monthly_ctc = salary_component.amount
                                        yearly_ctc = monthly_ctc * 12
                                    elif salary_component.rule_id.code == 'GRATUITY':
                                        monthly_gratuity = salary_component.amount
                                        yearly_gratuity = monthly_gratuity * 12
                                    elif salary_component.rule_id.code == 'PF_ER':
                                        monthly_pf_er = salary_component.amount
                                        yearly_pf_er = monthly_pf_er * 12
                                    elif salary_component.rule_id.code == 'PF_EE':
                                        monthly_pf_ee = salary_component.amount
                                        yearly_pf_ee = monthly_pf_ee * 12
                                    elif salary_component.rule_id.code == 'ESIC_ER':
                                        monthly_esic_er = salary_component.amount
                                        yearly_esic_er = monthly_esic_er * 12
                                    elif salary_component.rule_id.code == 'ESIC_EE':
                                        monthly_esic_ee = salary_component.amount
                                        yearly_esic_ee = monthly_esic_ee * 12
                                    elif salary_component.rule_id.code == 'HRA':
                                        monthly_hra = salary_component.amount
                                        yearly_hra = monthly_hra * 12
                                    elif salary_component.rule_id.code == 'Other':
                                        monthly_other = salary_component.amount
                                        yearly_other = monthly_other * 12
                                    elif salary_component.rule_id.code == 'GROSS':
                                        monthly_gross = salary_component.amount
                                        yearly_gross = monthly_gross * 12
                                    elif salary_component.rule_id.code == 'PT':
                                        monthly_pt = salary_component.amount
                                        yearly_pt = monthly_pt * 12
                                    elif salary_component.rule_id.code == 'NET':
                                        monthly_net = salary_component.amount
                                        yearly_net = monthly_net * 12

                                # Replace Values
                                if variable_content[0].strip() == 'salary_monthly_ctc':
                                    replace_variable_content = monthly_ctc
                                elif variable_content[0].strip() == 'salary_Yearly_ctc':
                                    replace_variable_content = yearly_ctc
                                elif variable_content[0].strip() == 'salary_monthly_gratuity':
                                    replace_variable_content = monthly_gratuity
                                elif variable_content[0].strip() == 'salary_Yearly_gratuity':
                                    replace_variable_content = yearly_gratuity
                                elif variable_content[0].strip() == 'salary_monthly_pf_er':
                                    replace_variable_content = monthly_pf_er
                                elif variable_content[0].strip() == 'salary_Yearly_pf_er':
                                    replace_variable_content = yearly_pf_er
                                elif variable_content[0].strip() == 'salary_monthly_esic_er':
                                    replace_variable_content = monthly_esic_er
                                elif variable_content[0].strip() == 'salary_Yearly_esic_er':
                                    replace_variable_content = yearly_esic_er
                                elif variable_content[0].strip() == 'salary_monthly_basic':
                                    replace_variable_content = monthly_basic
                                elif variable_content[0].strip() == 'salary_Yearly_basic':
                                    replace_variable_content = yearly_basic
                                elif variable_content[0].strip() == 'salary_monthly_hra':
                                    replace_variable_content = monthly_hra
                                elif variable_content[0].strip() == 'salary_Yearly_hra':
                                    replace_variable_content = yearly_hra
                                elif variable_content[0].strip() == 'salary_monthly_allowance':
                                    replace_variable_content = monthly_other
                                elif variable_content[0].strip() == 'salary_Yearly_allowance':
                                    replace_variable_content = yearly_other
                                elif variable_content[0].strip() == 'salary_monthly_gross':
                                    replace_variable_content = monthly_gross
                                elif variable_content[0].strip() == 'salary_Yearly_gross':
                                    replace_variable_content = yearly_gross
                                elif variable_content[0].strip() == 'salary_monthly_pt':
                                    replace_variable_content = monthly_pt
                                elif variable_content[0].strip() == 'salary_Yearly_pt':
                                    replace_variable_content = yearly_pt
                                elif variable_content[0].strip() == 'salary_monthly_pf_ee':
                                    replace_variable_content = monthly_pf_ee
                                elif variable_content[0].strip() == 'salary_Yearly_pf_ee':
                                    replace_variable_content = yearly_pf_ee
                                elif variable_content[0].strip() == 'salary_monthly_esic_ee':
                                    replace_variable_content = monthly_esic_ee
                                elif variable_content[0].strip() == 'salary_Yearly_esic_ee':
                                    replace_variable_content = yearly_esic_ee
                                elif variable_content[0].strip() == 'salary_monthly_net':
                                    replace_variable_content = monthly_net
                                elif variable_content[0].strip() == 'salary_Yearly_net':
                                    replace_variable_content = yearly_net
                                else:
                                    replace_variable_content = variable_content[0].strip()
                        else:
                            replace_variable_content = variable_content[0].strip()
                    elif len(variable_content) in [2, 3]:
                        outer_obj = getattr(employee_obj, variable_content[1].strip())
                        if len(variable_content) == 2:
                            replace_variable_content = outer_obj
                        else:
                            if '%d' in variable_content[2]:
                                if '%dth' in variable_content[2]:
                                    day = outer_obj.day
                                    day_suffix = ordinal_suffix_days.get(day,
                                                                         'th') if 10 <= day % 100 <= 20 else ordinal_suffix_days.get(
                                        day % 10, 'th')
                                    replace_variable_content = outer_obj.strftime(f'%d{day_suffix} %B, %Y')
                                else:
                                    replace_variable_content = outer_obj.strftime(variable_content[2].strip())
                            elif 'n2w' in variable_content[2]:
                                converting_number = outer_obj
                                words = num2words(round(converting_number), lang='en_IN')
                                replace_variable_content = words[0].capitalize() + words[1:]
                            else:
                                replace_variable_content = outer_obj
                    elif len(variable_content) in [4, 5]:
                        inner_obj = \
                            getattr(employee_obj, variable_content[1].strip()).sorted(key=lambda x: x.create_date,
                                                                                      reverse=True)[0]
                        repl_var_con_1 = getattr(inner_obj, variable_content[2].strip())
                        if len(variable_content) == 4:
                            replace_variable_content = repl_var_con_1
                        else:
                            if '%d' in variable_content[4]:
                                if '%dth' in variable_content[4]:
                                    day = repl_var_con_1.day
                                    day_suffix = ordinal_suffix_days.get(day,
                                                                         'th') if 10 <= day % 100 <= 20 else ordinal_suffix_days.get(
                                        day % 10, 'th')
                                    replace_variable_content = repl_var_con_1.strftime(f'%d{day_suffix} %B, %Y')
                                else:
                                    replace_variable_content = repl_var_con_1.strftime(variable_content[4].strip())
                            elif 'n2w' in variable_content[2]:
                                converting_number = repl_var_con_1
                                words = num2words(round(converting_number), lang='en_IN')
                                replace_variable_content = words[0].capitalize() + words[1:]
                            else:
                                replace_variable_content = repl_var_con_1
                    else:
                        replace_variable_content = getattr(employee_obj, variable_content[1].strip())

                    result += str(replace_variable_content)
                    i = end_index + 2
                else:
                    # If closing "}}" is not found, keep the original text
                    result += paragraph[i:]
                    break
            else:
                result += paragraph[i]
                i += 1

        return result

    def generate_employee_letter(self):
        employee_obj = self.env['hr.employee'].browse(self.env.context.get('employee_id'))
        appraisal_obj = self.env['employee.appraisal'].browse(self.env.context.get('appraisal_id'))
        if 'termination_letter' in self.env.context:
            letter_obj = self.env['hr.letters'].sudo().search(
                [('document_type_id', '=', self.env.ref('aspl_hr_employee.doc_type_termination').id)])
        else:
            letter_obj = self.letter

        # Select Letter Version
        letter_versions = letter_obj.letter_version_control_line_ids
        if len(letter_versions) == 0:
            raise ValidationError("Please add attachment in Hr Letter Section !!!")
        elif len(letter_versions) == 1:
            letter_doc_content = letter_versions.attachment_v_c
        else:
            if letter_obj.document_type_id.name == 'Offer Trainee':
                fetching_date = employee_obj.join_training_date
            elif (letter_obj.document_type_id.name == 'Offer Experience') or (
                    letter_obj.document_type_id.name == 'Appointment') or (
                    letter_obj.document_type_id.name == 'Appointment and Increment-Consultant'):
                fetching_date = employee_obj.join_date
            elif letter_obj.document_type_id.name == 'Confirmation':
                fetching_date = employee_obj.confirm_date
            elif letter_obj.document_type_id.name == 'Appraisal':
                fetching_date = appraisal_obj.appraisal_date
            elif letter_obj.document_type_id.name == 'Experience and Relieving':
                if employee_obj.leaving_date:
                    fetching_date = employee_obj.leaving_date
                else:
                    fetching_date = employee_obj.left_date
            elif letter_obj.document_type_id.name == 'Termination':
                fetching_date = employee_obj.left_date
            else:
                fetching_date = date.today()

            min_version_date = letter_versions.sorted(key=lambda x: x.create_date)[0].create_date
            max_version_date = letter_versions.sorted(key=lambda x: x.create_date, reverse=True)[0].create_date

            if min_version_date.date() >= fetching_date:
                letter_doc_content = letter_versions.sorted(key=lambda x: x.create_date)[0].attachment_v_c
            elif max_version_date.date() <= fetching_date:
                letter_doc_content = letter_versions.sorted(key=lambda x: x.create_date, reverse=True)[0].attachment_v_c
            else:
                updated_letter_version = letter_versions.filtered(lambda x: x.create_date.date() <= fetching_date)
                letter_doc_content = updated_letter_version.sorted(key=lambda x: x.create_date, reverse=True)[
                    0].attachment_v_c

        # Check Letter is Already in Attachment or not ...
        current_letter = employee_obj.employee_document_current_ids.filtered(
            lambda x: x.type_of_document == letter_obj.document_type_id)
        if not self.env.user.has_group('base.group_system'):
            if current_letter or appraisal_obj.document:
                if appraisal_obj.document:
                    if appraisal_obj.status == 'lock':
                        return {
                            'type': 'ir.actions.act_window',
                            'res_model': 'regeneration.letter.permission',
                            'view_type': 'form',
                            'view_mode': 'form',
                            'name': _('Document Exits.'),
                            'target': 'new',
                            'context': {'employee_id': employee_obj.id, 'letter_model': appraisal_obj._name,
                                        'letter_id': appraisal_obj.id},
                        }
                else:
                    if current_letter.type_of_document == 'Termination':
                        current_letter.status = 'unlock'
                    if current_letter.status == 'lock':
                        return {
                            'type': 'ir.actions.act_window',
                            'res_model': 'regeneration.letter.permission',
                            'view_type': 'form',
                            'view_mode': 'form',
                            'name': _('Document Exits.'),
                            'target': 'new',
                            'context': {'employee_id': employee_obj.id, 'letter_model': current_letter._name,
                                        'letter_id': current_letter.id},
                        }

        # Preparing a Doc File
        try:
            # Decode Base64 content
            file_content = base64.b64decode(letter_doc_content)

            # Load the Word document from the binary content
            doc = docx.Document(io.BytesIO(file_content))

            # Replace Content Values
            for paragraph in doc.paragraphs:
                if "{{ " in paragraph.text:
                    paragraph.text = self._change_content(paragraph.text, employee_obj)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if "{{ " in paragraph.text:
                                paragraph.text = self._change_content(paragraph.text, employee_obj)

            # Save the modified document
            modified_doc_content = io.BytesIO()
            doc.save(modified_doc_content)

            if current_letter or appraisal_obj.document:
                if appraisal_obj.document:
                    # Attachment Delete & Create New
                    attachment_id = self.env['ir.attachment'].sudo().search(
                        [('name', '=', appraisal_obj.document_name)])
                    attachment_id.unlink()

                    vals = {
                        'name': f'{letter_obj.name} - {employee_obj.name}.docx',
                        'datas': base64.b64encode(modified_doc_content.getvalue()),
                        'store_fname': f'{letter_obj.name} - {employee_obj.name}.docx',
                        'res_model': 'hr.employee',
                        'res_id': self.id,
                    }
                    attachment_id = self.env['ir.attachment'].create(vals)

                    # Update Appraisal Latter
                    appraisal_obj.document_name = attachment_id.name
                    appraisal_obj.document = attachment_id.datas
                    appraisal_obj.status = 'lock'

                    # Add Log Note
                    employee_obj.message_post(
                        body="Appraisal Document has been Updated. - %s" % appraisal_obj.appraisal_date.strftime(
                            "%d/%m/%Y"))
                else:
                    # Attachment Delete & Create New
                    attachment_id = self.env['ir.attachment'].sudo().search(
                        [('name', '=', current_letter.document_name)])
                    attachment_id.unlink()

                    vals = {
                        'name': f'{letter_obj.name} - {employee_obj.name}.docx',
                        'datas': base64.b64encode(modified_doc_content.getvalue()),
                        'store_fname': f'{letter_obj.name} - {employee_obj.name}.docx',
                        'res_model': 'hr.employee',
                        'res_id': self.id,
                    }
                    attachment_id = self.env['ir.attachment'].create(vals)

                    # Document Update
                    current_letter.document_name = attachment_id.name
                    current_letter.attached_date = date.today()
                    current_letter.document = attachment_id.datas
                    current_letter.status = 'lock'

                    # Add Log Note
                    employee_obj.message_post(
                        body="%s Document has been Updated." % current_letter.type_of_document.name)
            else:
                # Attachment Create
                vals = {
                    'name': f'{letter_obj.name} - {employee_obj.name}.docx',
                    'datas': base64.b64encode(modified_doc_content.getvalue()),
                    'store_fname': f'{letter_obj.name} - {employee_obj.name}.docx',
                    'res_model': 'hr.employee',
                    'res_id': self.id,
                }
                attachment_id = self.env['ir.attachment'].create(vals)

                if appraisal_obj:
                    # Add Appraisal Latter In Current Id
                    appraisal_obj.document_name = attachment_id.name
                    appraisal_obj.document = attachment_id.datas
                    appraisal_obj.status = 'lock'

                    # Add Log Note
                    employee_obj.message_post(
                        body="Appraisal Document created - %s" % appraisal_obj.appraisal_date.strftime("%d/%m/%Y"))
                else:
                    # Current Employee Document Create
                    vals = {
                        'employee_id': employee_obj.id,
                        'document_name': attachment_id.name,
                        'type': "current",
                        'type_of_document': letter_obj.document_type_id.id,
                        'report_type': "word",
                        'attached_date': date.today(),
                        'document': attachment_id.datas,
                        'status': 'lock',
                    }
                    document_id = self.env['employee.document'].create(vals)

                    # Add Log Note
                    employee_obj.message_post(body="%s Document created." % document_id.type_of_document.name)

            if 'termination_letter' in self.env.context:
                if not current_letter:
                    document_id.status = 'unlock'
                else:
                    current_letter.status = 'unlock'
                return attachment_id
            else:
                # Download Attachment File
                download_url = '/web/content/' + \
                               str(attachment_id.id) + '?download=True'
                base_url = self.env['ir.config_parameter'].sudo().get_param('web.base.url')

                return {
                    "type": "ir.actions.act_url",
                    "url": str(base_url) + str(download_url),
                    "target": "new"
                }

        except Exception as e:
            # Raise Error
            raise ValidationError(f"Error: {e}")
